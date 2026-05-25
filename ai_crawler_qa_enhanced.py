
import os
import re
import json
import time
import hashlib
from io import BytesIO
from collections import deque
from urllib.parse import urljoin, urlparse, urlunparse, parse_qsl, urlencode

import requests
from bs4 import BeautifulSoup
from tqdm import tqdm
from ddgs import DDGS
from openai import OpenAI
import fitz
import docx

try:
    import openpyxl
except Exception:
    openpyxl = None


class EnhancedAICampusQACrawler:
    """
    Enhanced SUSTech campus QA crawler.

    Features:
    - Append to existing JSONL files without overwriting.
    - Deduplicate by normalized URL, content hash, and normalized question hash.
    - Use high-value seed URLs + web search + 1~2 depth internal link discovery.
    - Parse HTML/PDF/DOCX/XLSX.
    - Save skipped/failed URLs to sustech_failed_urls.jsonl.
    - Generate QA pairs with evidence snippets for better RAG explainability.
    """

    def __init__(
        self,
        user_requirement,
        allowed_domains=None,
        seed_urls=None,
        extra_keywords=None,
        max_search_keywords=60,
        max_urls_per_keyword=20,
        max_pages=300,
        max_depth=2,
        max_links_per_page=80,
        max_qa_per_chunk=8,
        chunk_size=2200,
        overlap=300,
        output_file="sustech_qa_pairs.jsonl",
        raw_output_file="sustech_raw_documents.jsonl",
        failed_output_file="sustech_failed_urls.jsonl",
        model="deepseek-chat",
        delay=0.6,
        skip_existing_urls=True,
        skip_existing_questions=True,
        skip_duplicate_content=True,
        force_recrawl=False,
        use_ai_url_filter=False,
        use_ai_content_filter=True,
        min_content_length=120,
    ):
        self.user_requirement = user_requirement
        self.allowed_domains = allowed_domains or []
        self.seed_urls = seed_urls or []
        self.extra_keywords = extra_keywords or []
        self.max_search_keywords = max_search_keywords
        self.max_urls_per_keyword = max_urls_per_keyword
        self.max_pages = max_pages
        self.max_depth = max_depth
        self.max_links_per_page = max_links_per_page
        self.max_qa_per_chunk = max_qa_per_chunk
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.output_file = output_file
        self.raw_output_file = raw_output_file
        self.failed_output_file = failed_output_file
        self.model = model
        self.delay = delay
        self.skip_existing_urls = skip_existing_urls
        self.skip_existing_questions = skip_existing_questions
        self.skip_duplicate_content = skip_duplicate_content
        self.force_recrawl = force_recrawl
        self.use_ai_url_filter = use_ai_url_filter
        self.use_ai_content_filter = use_ai_content_filter
        self.min_content_length = min_content_length

        self.client = OpenAI(api_key=os.getenv("DEEPSEEK_API_KEY"), base_url="https://api.deepseek.com")
        self.session = requests.Session()
        self.session.headers.update({
            "User-Agent": "Mozilla/5.0 (compatible; Educational-RAG-Crawler/2.0)"
        })

        self.visited_urls = set()
        self.saved_questions = set()
        self.saved_doc_hashes = set()
        self.failed_keys = set()

        self.important_keywords = [
            "校园卡", "学生证", "一卡通", "挂失", "补办", "充值",
            "宿舍", "公寓", "住宿", "退宿", "入住", "门禁",
            "选课", "加课", "退课", "课程", "教务", "学分", "成绩单", "在读证明",
            "休学", "复学", "退学", "转学", "学籍", "请假",
            "图书馆", "开馆", "借阅", "讨论间", "空间预约", "读者证", "数据库", "校外访问",
            "食堂", "餐厅", "校巴", "总务", "校园服务",
            "校医院", "医保", "就诊", "体检", "心理",
            "校园网", "网络", "VPN", "邮箱", "Blackboard", "雨课堂",
            "办事指南", "学生事务", "网上服务大厅", "服务大厅",
            "奖学金", "助学金", "勤工助学", "资助", "荣誉",
            "新生", "学生手册", "管理规定", "规章制度", "通知", "FAQ", "常见问题",
        ]

    def clean_json_text(self, text):
        text = (text or "").strip()
        text = re.sub(r"^```json", "", text).strip()
        text = re.sub(r"^```", "", text).strip()
        text = re.sub(r"```$", "", text).strip()
        if not text.startswith("{"):
            s, e = text.find("{"), text.rfind("}")
            if s >= 0 and e > s:
                text = text[s:e + 1]
        return text

    def ask_ai_json(self, prompt, temperature=0.2, retries=2):
        for i in range(retries + 1):
            try:
                resp = self.client.chat.completions.create(
                    model=self.model,
                    messages=[
                        {"role": "system", "content": "你是严谨的数据处理助手，只输出合法 JSON，不要 Markdown。"},
                        {"role": "user", "content": prompt},
                    ],
                    temperature=temperature,
                )
                return json.loads(self.clean_json_text(resp.choices[0].message.content))
            except Exception as e:
                print(f"[AI JSON failed] {i + 1}/{retries + 1}: {e}")
                time.sleep(1 + i)
        return {}

    def save_jsonl(self, path, record):
        with open(path, "a", encoding="utf-8") as f:
            f.write(json.dumps(record, ensure_ascii=False) + "\n")

    def make_id(self, text):
        return hashlib.md5(str(text).encode("utf-8")).hexdigest()

    def normalize_url(self, url):
        if not url:
            return ""
        url = str(url).strip().replace("\\", "/").split("#")[0].strip()
        p = urlparse(url)
        if not p.scheme:
            return url.rstrip("/")
        keep = []
        for k, v in parse_qsl(p.query, keep_blank_values=True):
            lk = k.lower()
            if lk.startswith("utm_") or lk in {"from", "spm", "tdsourcetag"}:
                continue
            keep.append((k, v))
        q = urlencode(keep, doseq=True)
        out = urlunparse((p.scheme.lower(), p.netloc.lower(), p.path, "", q, ""))
        return out[:-1] if out.endswith("/") and p.path != "/" else out

    def normalize_question(self, q):
        return re.sub(r"\s+", "", str(q).strip().replace("？", "?"))

    def get_content_hash(self, content):
        content = re.sub(r"\s+", "", str(content or "").strip())
        return self.make_id(content[:12000])

    def domain_allowed(self, url):
        if not self.allowed_domains:
            return True
        host = urlparse(url).netloc.lower()
        full = url.lower()
        return any(host == d.lower() or host.endswith("." + d.lower()) or d.lower() in full for d in self.allowed_domains)

    def log_failed(self, url, reason, **kwargs):
        url = self.normalize_url(url)
        key = f"{url}|{reason}"
        if not url or key in self.failed_keys:
            return
        self.failed_keys.add(key)
        rec = {"url": url, "reason": reason, "time": time.strftime("%Y-%m-%d %H:%M:%S")}
        rec.update(kwargs)
        self.save_jsonl(self.failed_output_file, rec)

    def load_jsonl(self, path):
        if not os.path.exists(path):
            return []
        out = []
        with open(path, "r", encoding="utf-8") as f:
            for i, line in enumerate(f, 1):
                line = line.strip()
                if not line:
                    continue
                try:
                    out.append(json.loads(line))
                except Exception:
                    print(f"[WARN] bad json line {i}: {path}")
        return out

    def load_existing_state(self):
        raw = self.load_jsonl(self.raw_output_file)
        for r in raw:
            url = self.normalize_url(r.get("url", ""))
            if url:
                self.visited_urls.add(url)
            if r.get("content_hash"):
                self.saved_doc_hashes.add(r["content_hash"])
            elif r.get("content"):
                self.saved_doc_hashes.add(self.get_content_hash(r["content"]))

        qa = self.load_jsonl(self.output_file)
        for r in qa:
            if r.get("question"):
                self.saved_questions.add(r.get("id") or self.make_id(self.normalize_question(r["question"])))

        failed = self.load_jsonl(self.failed_output_file)
        for r in failed:
            if r.get("url") and r.get("reason"):
                self.failed_keys.add(f"{self.normalize_url(r['url'])}|{r['reason']}")

        print(f"[Loaded] raw={len(raw)}, qa={len(qa)}, failed={len(failed)}")
        print(f"[Loaded] urls={len(self.visited_urls)}, content_hash={len(self.saved_doc_hashes)}, questions={len(self.saved_questions)}")

    def is_useful_link(self, url, anchor="", title=""):
        if not self.domain_allowed(url):
            return False
        u = url.lower()
        if any(x in u for x in ["javascript:", "mailto:", "tel:", "/login", "logout"]):
            return False
        if any(u.endswith(x) for x in [".jpg", ".png", ".gif", ".svg", ".webp", ".mp4", ".mp3"]):
            return False
        if any(u.endswith(x) for x in [".pdf", ".docx", ".doc", ".xlsx", ".xls"]):
            return True
        text = f"{url} {anchor} {title}".lower()
        if any(k.lower() in text for k in self.important_keywords):
            return True
        return any(m in u for m in ["list", "page", "student", "service", "guide", "notice", "faq", "cjwt", "jyzd"])

    def generate_search_keywords(self):
        prompt = f"""
用户想构建南方科技大学校园知识库 QA 数据集。需求：
{self.user_requirement}
请生成最多 {self.max_search_keywords} 个搜索关键词，覆盖学生事务、课程、成绩单、宿舍、图书馆、校园卡、校园网、Blackboard、校医院、食堂、校巴、学生手册、新生攻略。
只返回 JSON：{{"keywords":["..."]}}
"""
        result = self.ask_ai_json(prompt)
        keywords = result.get("keywords", []) if isinstance(result.get("keywords", []), list) else []
        defaults = [
            "site:sscsop.sustech.edu.cn 南方科技大学 学生事务办理指南",
            "site:sscsop.sustech.edu.cn 南方科技大学 成绩单 打印",
            "site:sscsop.sustech.edu.cn 南方科技大学 在读证明 打印",
            "site:sscsop.sustech.edu.cn 南方科技大学 选课咨询",
            "site:sscsop.sustech.edu.cn 南方科技大学 学分认定",
            "site:sscsop.sustech.edu.cn 南方科技大学 休学 复学",
            "site:sscsop.sustech.edu.cn 南方科技大学 退课 办理",
            "site:sscsop.sustech.edu.cn 南方科技大学 学生请假",
            "site:sscsop.sustech.edu.cn 南方科技大学 勤工助学",
            "site:lib.sustech.edu.cn 南方科技大学 图书馆 开馆时间",
            "site:lib.sustech.edu.cn 南方科技大学 图书馆 借阅制度",
            "site:lib.sustech.edu.cn 南方科技大学 图书馆 讨论间 预约",
            "site:lib.sustech.edu.cn 南方科技大学 图书馆 数据库 校外访问",
            "site:gao.sustech.edu.cn 南方科技大学 总务 服务手册 PDF",
            "site:gao.sustech.edu.cn 南方科技大学 学生公寓 管理细则",
            "site:gao.sustech.edu.cn 南方科技大学 食堂 餐厅",
            "site:www.sustech.edu.cn 南方科技大学 学生常用系统 Blackboard",
            "site:mirrors.sustech.edu.cn 南方科技大学 学生手册 PDF",
            "site:mirrors.sustech.edu.cn 南方科技大学 新生攻略 PDF",
            "南方科技大学 校园卡 挂失 补办",
            "南方科技大学 校园网 VPN 邮箱 使用指南",
            "南方科技大学 Blackboard 使用指南",
            "南方科技大学 校医院 就诊 医保",
            "南方科技大学 校巴 时刻表",
            "南方科技大学 校历",
        ]
        for kw in self.extra_keywords + defaults:
            if kw not in keywords:
                keywords.append(kw)
        return keywords[:self.max_search_keywords]

    def search_web(self, keyword):
        out = []
        try:
            with DDGS() as ddgs:
                for r in ddgs.text(keyword, max_results=self.max_urls_per_keyword):
                    url = self.normalize_url(r.get("href", ""))
                    if url and self.domain_allowed(url):
                        out.append({"url": url, "title": r.get("title", ""), "snippet": r.get("body", ""), "source_type": "search", "depth": 0})
        except Exception as e:
            print(f"[Search failed] {keyword}: {e}")
        return out

    def fetch_content(self, url):
        url = self.normalize_url(url)
        try:
            resp = self.session.get(url, timeout=35, allow_redirects=True)
            resp.raise_for_status()
            final_url = self.normalize_url(resp.url)
            ctype = resp.headers.get("Content-Type", "").lower()
            u = final_url.lower()

            if "text/html" in ctype or u.endswith((".html", ".htm", ".psp")) or "." not in urlparse(final_url).path.split("/")[-1]:
                resp.encoding = resp.apparent_encoding or "utf-8"
                title, content, links = self.extract_html(resp.text, final_url)
                return {"url": final_url, "title": title, "content": content, "links": links, "content_type": ctype or "text/html"}

            if "pdf" in ctype or u.endswith(".pdf"):
                title, content = self.extract_pdf(resp.content, final_url)
                return {"url": final_url, "title": title, "content": content, "links": [], "content_type": ctype or "pdf"}

            if "wordprocessingml" in ctype or u.endswith(".docx"):
                title, content = self.extract_docx(resp.content, final_url)
                return {"url": final_url, "title": title, "content": content, "links": [], "content_type": ctype or "docx"}

            if "spreadsheetml" in ctype or u.endswith(".xlsx"):
                title, content = self.extract_xlsx(resp.content, final_url)
                return {"url": final_url, "title": title, "content": content, "links": [], "content_type": ctype or "xlsx"}

            if u.endswith(".doc"):
                self.log_failed(final_url, "unsupported_doc", content_type=ctype)
                return None
            if u.endswith(".xls"):
                self.log_failed(final_url, "unsupported_xls", content_type=ctype)
                return None

            self.log_failed(final_url, "unsupported_content_type", content_type=ctype)
            return None
        except Exception as e:
            self.log_failed(url, "fetch_failed", error=str(e))
            return None

    def extract_html(self, html, base_url):
        soup = BeautifulSoup(html, "lxml")
        title = soup.title.get_text(strip=True) if soup.title else ""

        links, seen = [], set()
        for a in soup.find_all("a", href=True):
            child = self.normalize_url(urljoin(base_url, a["href"]))
            anchor = a.get_text(" ", strip=True)
            if child and child not in seen and child != base_url and self.is_useful_link(child, anchor, title):
                seen.add(child)
                links.append({"url": child, "anchor": anchor, "from_url": base_url})

        for tag in soup(["script", "style", "nav", "footer", "header", "aside", "form", "noscript", "iframe"]):
            tag.decompose()

        main = (
            soup.find("main")
            or soup.find("article")
            or soup.find("div", class_=re.compile("content|main|article|entry|detail|wp_articlecontent", re.I))
            or soup.find("div", id=re.compile("content|main|article|detail", re.I))
            or soup.body
            or soup
        )

        table_lines = []
        for table in main.find_all("table"):
            for row in table.find_all("tr"):
                cells = [c.get_text(" ", strip=True) for c in row.find_all(["td", "th"])]
                cells = [re.sub(r"\s+", " ", c).strip() for c in cells if c.strip()]
                if cells:
                    table_lines.append(" | ".join(cells))

        lines = []
        for line in main.get_text("\n", strip=True).split("\n"):
            line = re.sub(r"\s+", " ", line).strip()
            if len(line) >= 3:
                lines.append(line)

        dedup, last = [], None
        for line in lines:
            if line != last:
                dedup.append(line)
            last = line

        content = "\n".join(dedup)
        if table_lines:
            content += "\n\n[表格内容]\n" + "\n".join(table_lines)
        return title, content, links[:self.max_links_per_page]

    def extract_pdf(self, data, url):
        try:
            pdf = fitz.open(stream=data, filetype="pdf")
            chunks = []
            for i, page in enumerate(pdf):
                text = page.get_text("text").strip()
                if text:
                    chunks.append(f"第 {i + 1} 页\n{text}")
            return url.split("/")[-1] or "PDF", "\n\n".join(chunks)
        except Exception as e:
            self.log_failed(url, "pdf_parse_failed", error=str(e))
            return None, None

    def extract_docx(self, data, url):
        try:
            d = docx.Document(BytesIO(data))
            lines = [p.text.strip() for p in d.paragraphs if p.text.strip()]
            for table in d.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return url.split("/")[-1] or "DOCX", "\n".join(lines)
        except Exception as e:
            self.log_failed(url, "docx_parse_failed", error=str(e))
            return None, None

    def extract_xlsx(self, data, url):
        if openpyxl is None:
            self.log_failed(url, "xlsx_parse_failed", error="openpyxl not installed")
            return None, None
        try:
            wb = openpyxl.load_workbook(BytesIO(data), read_only=True, data_only=True)
            lines = []
            for ws in wb.worksheets[:10]:
                lines.append(f"[工作表] {ws.title}")
                for row in ws.iter_rows(max_row=500, values_only=True):
                    cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                    if cells:
                        lines.append(" | ".join(cells))
            return url.split("/")[-1] or "XLSX", "\n".join(lines)
        except Exception as e:
            self.log_failed(url, "xlsx_parse_failed", error=str(e))
            return None, None

    def judge_content_value(self, title, url, content):
        prompt = f"""
请判断下面资料是否适合生成“南方科技大学校园知识库 QA 对”。

用户需求：
{self.user_requirement}

标题：{title}
URL：{url}
正文节选：
{content[:3500]}

如果资料包含校园服务、学生事务、教学事务、图书馆、住宿、校园网、Blackboard、校医院、校巴、学生手册、办事流程、规章制度、FAQ、表格说明等，返回 useful=true。
如果只是新闻宣传、人物介绍、科研成果、纯导航且没有具体服务信息，返回 useful=false。
category 使用英文：course/dormitory/library/campus_card/network/medical/student_affairs/logistics/scholarship/new_student/rules/other。

只返回 JSON：
{{"useful": true, "category": "student_affairs", "reason": "简短原因"}}
"""
        result = self.ask_ai_json(prompt)
        if "useful" not in result:
            result["useful"] = False
        return result

    def split_text(self, text):
        text = (text or "").strip()
        out, start = [], 0
        while start < len(text):
            end = min(start + self.chunk_size, len(text))
            chunk = text[start:end].strip()
            if len(chunk) >= 250:
                out.append(chunk)
            if end >= len(text):
                break
            start = max(0, end - self.overlap)
        return out

    def generate_qa_pairs(self, title, url, content, category):
        prompt = f"""
请根据下面的南方科技大学校园资料生成高质量 QA 对。

资料标题：{title}
资料 URL：{url}
资料分类：{category}
资料正文：
{content}

要求：
1. 只能根据资料正文生成，不允许编造。
2. 问题像学生真实会问的问题，覆盖“怎么/哪里/什么时候/需要什么材料/流程”等表达。
3. 答案保留关键时间、地点、入口、材料、流程、联系方式。
4. 每条 QA 带 source_title、source_url、category、evidence。
5. evidence 必须是正文依据片段，30-250 字。
6. 对时效性信息，在答案末尾补充“具体安排以学校最新官方通知为准”。
7. 最多生成 {self.max_qa_per_chunk} 条；无价值返回空数组。
只返回 JSON：
{{"qa_pairs":[{{"question":"问题","answer":"答案","category":"{category}","source_title":"资料标题","source_url":"资料URL","evidence":"原文依据片段"}}]}}
"""
        result = self.ask_ai_json(prompt)
        pairs = result.get("qa_pairs", [])
        if not isinstance(pairs, list):
            return []
        cleaned = []
        for qa in pairs:
            if not isinstance(qa, dict):
                continue
            q, a = str(qa.get("question", "")).strip(), str(qa.get("answer", "")).strip()
            if len(q) < 5 or len(a) < 5:
                continue
            qa["question"], qa["answer"] = q, a
            qa["category"] = qa.get("category", category or "other")
            qa["source_title"] = qa.get("source_title", title)
            qa["source_url"] = qa.get("source_url", url)
            qa["evidence"] = str(qa.get("evidence", "")).strip()[:500] or content[:250]
            cleaned.append(qa)
        return cleaned

    def build_initial_queue(self):
        queue, queued = deque(), set()

        def push(url, title="", snippet="", source_type="seed", depth=0):
            url = self.normalize_url(url)
            if url and url not in queued and self.domain_allowed(url):
                queued.add(url)
                queue.append({"url": url, "title": title, "snippet": snippet, "source_type": source_type, "depth": depth})

        for u in self.seed_urls:
            push(u, title="Seed URL", source_type="seed", depth=0)

        print("\nStep 1: 生成搜索关键词...")
        keywords = self.generate_search_keywords()
        for kw in keywords:
            print("-", kw)

        print("\nStep 2: 搜索网页...")
        for kw in keywords:
            print(f"\n搜索关键词：{kw}")
            for item in self.search_web(kw):
                push(item["url"], title=item.get("title", ""), snippet=item.get("snippet", ""), source_type="search", depth=0)
                time.sleep(0.1)

        print(f"\n[INFO] 初始候选 URL 数量：{len(queue)}")
        return queue, queued

    def process_document(self, item, fetched):
        url = fetched["url"]
        title = fetched.get("title") or item.get("title") or url
        content = fetched.get("content") or ""

        if len(content) < self.min_content_length:
            self.log_failed(url, "content_too_short", title=title, length=len(content))
            return 0, False

        chash = self.get_content_hash(content)
        if self.skip_duplicate_content and chash in self.saved_doc_hashes:
            self.log_failed(url, "duplicate_content", title=title, content_hash=chash)
            return 0, False

        category, reason = item.get("ai_category", "other"), item.get("ai_reason", "")
        if self.use_ai_content_filter:
            judge = self.judge_content_value(title, url, content)
            if judge.get("useful") is not True:
                self.log_failed(url, "ai_content_irrelevant", title=title, reason_detail=judge.get("reason", ""))
                return 0, False
            category = judge.get("category", category or "other")
            reason = judge.get("reason", reason)

        self.saved_doc_hashes.add(chash)
        self.save_jsonl(self.raw_output_file, {
            "doc_id": self.make_id(url),
            "url": url,
            "title": title,
            "content": content,
            "content_hash": chash,
            "content_length": len(content),
            "content_type": fetched.get("content_type", ""),
            "ai_category": category,
            "ai_reason": reason,
            "crawl_source": item.get("source_type", "unknown"),
            "crawl_depth": item.get("depth", 0),
            "crawl_time": time.strftime("%Y-%m-%d %H:%M:%S"),
        })

        total = 0
        for chunk_id, chunk in enumerate(self.split_text(content), 1):
            for qa in self.generate_qa_pairs(f"{title} - chunk {chunk_id}", url, chunk, category):
                qid = self.make_id(self.normalize_question(qa["question"]))
                if self.skip_existing_questions and qid in self.saved_questions:
                    continue
                self.saved_questions.add(qid)
                qa["id"] = qid
                qa["crawl_source"] = "enhanced_deepseek_ai_crawler"
                qa["created_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
                qa["chunk_id"] = chunk_id
                qa["source_doc_id"] = self.make_id(url)
                qa["source_url"] = qa.get("source_url", url)
                qa["source_title"] = qa.get("source_title", title)
                self.save_jsonl(self.output_file, qa)
                total += 1
            time.sleep(0.25)
        return total, True

    def run(self):
        self.load_existing_state()
        queue, queued = self.build_initial_queue()
        print("\nStep 3: 抓取网页/文件并递归发现站内链接。")

        processed = useful_docs = new_qa = skipped_existing = 0
        pbar = tqdm(total=self.max_pages, desc="Crawling")

        while queue and processed < self.max_pages:
            item = queue.popleft()
            url = self.normalize_url(item["url"])
            depth = int(item.get("depth", 0))

            if self.skip_existing_urls and not self.force_recrawl and url in self.visited_urls:
                skipped_existing += 1
                continue

            self.visited_urls.add(url)
            fetched = self.fetch_content(url)
            if not fetched:
                continue

            final_url = self.normalize_url(fetched.get("url", url))
            if depth < self.max_depth:
                for link in fetched.get("links", []):
                    child = self.normalize_url(link.get("url", ""))
                    if child and child not in queued and self.domain_allowed(child):
                        queued.add(child)
                        queue.append({
                            "url": child,
                            "title": link.get("anchor", ""),
                            "snippet": f"discovered from {final_url}",
                            "source_type": "site_discovery",
                            "depth": depth + 1,
                        })

            qa_count, useful = self.process_document(item, fetched)
            new_qa += qa_count
            useful_docs += int(useful)
            processed += 1
            pbar.update(1)
            time.sleep(self.delay)

        pbar.close()
        print("\n完成。")
        print(f"新处理页面数量：{processed}")
        print(f"有效文档数量：{useful_docs}")
        print(f"跳过历史 URL 数量：{skipped_existing}")
        print(f"新增 QA 对数量：{new_qa}")
        print(f"已发现 URL 数量：{len(queued)}")
        print(f"原始网页数据追加保存到：{self.raw_output_file}")
        print(f"QA 对追加保存到：{self.output_file}")
        print(f"失败/跳过日志保存到：{self.failed_output_file}")


if __name__ == "__main__":
    requirement = """
我正在做南方科技大学校园知识库项目。
重点关注：校园卡、宿舍、选课、退课、成绩单、在读证明、图书馆、食堂、校医院、学生事务、Blackboard、校园网、学生手册、新生攻略、办事指南、规章制度、奖助学金、校巴等。
要求：优先使用南方科技大学官网、学生事务办理指南、图书馆、教学工作部、总务与空间办公室、官方通知、PDF/DOCX/XLSX 附件等公开信息；答案必须基于原文，不允许编造。
"""

    HIGH_VALUE_SEED_URLS = [
        "https://www.sustech.edu.cn/",
        "https://www.sustech.edu.cn/zh/students.html",
        "https://www.sustech.edu.cn/zh/common-systems.html",
        "https://sscsop.sustech.edu.cn/students/main.psp",
        "https://sscsop.sustech.edu.cn/2744/list2.htm",
        "https://sscsop.sustech.edu.cn/2747/list.htm",
        "https://sscsop.sustech.edu.cn/2021/0803/c1116a5402/page.htm",
        "https://sscsop.sustech.edu.cn/2021/0803/c1116a5397/page.htm",
        "https://sscsop.sustech.edu.cn/2023/1108/c1116a7320/page.htm",
        "https://sscsop.sustech.edu.cn/2023/1108/c1117a7300/page.htm",
        "https://sscsop.sustech.edu.cn/2021/0803/c1116a5398/page.htm",
        "https://sscsop.sustech.edu.cn/2023/1108/c1119a7340/page.htm",
        "https://lib.sustech.edu.cn/",
        "https://lib.sustech.edu.cn/kgsj_77/list.htm",
        "https://lib.sustech.edu.cn/jyzd/list.htm",
        "https://lib.sustech.edu.cn/kjyy/list.htm",
        "https://lib.sustech.edu.cn/224/list.htm",
        "https://lib.sustech.edu.cn/sjk/list.htm",
        "https://lib.sustech.edu.cn/cjwt/list.htm",
        "https://lib.sustech.edu.cn/1402/list53.htm",
        "https://lib.sustech.edu.cn/szmtkj/list.htm",
        "https://gao.sustech.edu.cn/",
        "https://gao.sustech.edu.cn/space/housing.html?lang=zh-cn",
        "https://gao.sustech.edu.cn/uploads/202402/23102621_74058.pdf",
        "https://mirrors.sustech.edu.cn/site/sustech-online/",
        "https://mirrors.sustech.edu.cn/git/sustech-online/sustech-online-ng/-/tree/master/docs/service?ref_type=heads",
        "https://mirrors.sustech.edu.cn/site/sustech-online/documents/manual/%E5%8D%97%E6%96%B9%E7%A7%91%E6%8A%80%E5%A4%A7%E5%AD%A6%E5%AD%A6%E7%94%9F%E6%89%8B%E5%86%8C2022-%E6%9C%AC%E7%A7%91.pdf",
        "https://mirrors.sustech.edu.cn/site/sustech-online/documents/college/zhicheng/2022%E6%96%B0%E7%94%9F%E6%94%BB%E7%95%A5by%E8%87%B4%E8%AF%9A%E4%B9%A6%E9%99%A2.pdf",
        "https://mirrors.sustech.edu.cn/site/sustech-online/documents/college/zhiren/%E8%87%B4%E4%BB%81%E4%B9%A6%E9%99%A2%E6%96%B0%E7%94%9F%E6%94%BB%E7%95%A5%E6%89%8B%E5%86%8C2022.pdf",
        "https://ehall.sustech.edu.cn/new/index.html",
    ]

    crawler = EnhancedAICampusQACrawler(
        user_requirement=requirement,
        allowed_domains=[
            "sustech.edu.cn",
            "www.sustech.edu.cn",
            "sscsop.sustech.edu.cn",
            "osa.sustech.edu.cn",
            "lib.sustech.edu.cn",
            "teaching.sustech.edu.cn",
            "tao.sustech.edu.cn",
            "gao.sustech.edu.cn",
            "student.sustech.edu.cn",
            "mirrors.sustech.edu.cn",
            "ehall.sustech.edu.cn",
            "tis.sustech.edu.cn",
        ],
        seed_urls=HIGH_VALUE_SEED_URLS,
        extra_keywords=[
            "site:sscsop.sustech.edu.cn 南方科技大学 退课办理",
            "site:sscsop.sustech.edu.cn 南方科技大学 在读证明打印封装",
            "site:sscsop.sustech.edu.cn 南方科技大学 学生请假管理",
            "site:sscsop.sustech.edu.cn 南方科技大学 退学 转学",
            "site:sscsop.sustech.edu.cn 南方科技大学 奖学金 助学金",
            "site:sscsop.sustech.edu.cn 南方科技大学 户籍 档案",
            "site:lib.sustech.edu.cn 南方科技大学 数据库访问问题 校外访问",
            "site:gao.sustech.edu.cn 南方科技大学 餐厅 食堂 校巴",
            "site:mirrors.sustech.edu.cn sustech-online 校园卡 学生证",
            "site:mirrors.sustech.edu.cn sustech-online 校园网 Blackboard",
        ],
        max_search_keywords=60,
        max_urls_per_keyword=20,
        max_pages=300,
        max_depth=2,
        max_links_per_page=80,
        max_qa_per_chunk=8,
        output_file="sustech_qa_pairs.jsonl",
        raw_output_file="sustech_raw_documents.jsonl",
        failed_output_file="sustech_failed_urls.jsonl",
        model="deepseek-chat",
        delay=0.6,
        skip_existing_urls=True,
        skip_existing_questions=True,
        skip_duplicate_content=True,
        force_recrawl=False,
        use_ai_url_filter=False,
        use_ai_content_filter=True,
        min_content_length=120,
    )
    crawler.run()
